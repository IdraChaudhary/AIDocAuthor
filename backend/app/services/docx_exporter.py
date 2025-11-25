from docx import Document
def build_docx(project, filename):
    doc = Document()
    doc.add_heading(project.title, level=1)
    for s in sorted(project.sections, key=lambda x: x.order):
        doc.add_heading(s.title, level=2)
        doc.add_paragraph(s.content or '')
    doc.save(filename)
